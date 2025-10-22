"""
Конвертер Markdown в Word
Поддерживает таблицы, заголовки, списки, форматирование текста
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    """Класс для конвертации Markdown в Word"""

    def __init__(self):
        self.doc = None

    def create_document(self):
        """Создает новый документ Word"""
        self.doc = Document()
        # Устанавливаем стандартный стиль
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет

    def add_table_borders(self, table):
        """Добавляет границы к таблице"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Создаем элемент границ
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def parse_inline_formatting(self, text):
        """Парсит встроенное форматирование (жирный, курсив, код)"""
        # Возвращает список кортежей (текст, форматирование)
        # форматирование: {'bold': bool, 'italic': bool, 'code': bool}
        parts = []

        # Регулярное выражение для поиска форматирования
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|___.*?___|__.*?__|_.*?_)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Добавляем текст до совпадения
            if match.start() > last_end:
                parts.append((text[last_end:match.start()], {}))

            matched_text = match.group()
            formatting = {}
            clean_text = matched_text

            # Проверяем тип форматирования
            if matched_text.startswith('***') or matched_text.startswith('___'):
                formatting = {'bold': True, 'italic': True}
                clean_text = matched_text[3:-3]
            elif matched_text.startswith('**') or matched_text.startswith('__'):
                formatting = {'bold': True}
                clean_text = matched_text[2:-2]
            elif matched_text.startswith('*') or matched_text.startswith('_'):
                formatting = {'italic': True}
                clean_text = matched_text[1:-1]
            elif matched_text.startswith('`'):
                formatting = {'code': True}
                clean_text = matched_text[1:-1]

            parts.append((clean_text, formatting))
            last_end = match.end()

        # Добавляем оставшийся текст
        if last_end < len(text):
            parts.append((text[last_end:], {}))

        return parts if parts else [(text, {})]

    def add_formatted_text(self, paragraph, text):
        """Добавляет текст с форматированием в параграф"""
        parts = self.parse_inline_formatting(text)

        for part_text, formatting in parts:
            run = paragraph.add_run(part_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True
            if formatting.get('code'):
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    def process_table(self, lines, start_idx):
        """Обрабатывает таблицу из markdown"""
        table_lines = []
        idx = start_idx

        # Собираем все строки таблицы
        while idx < len(lines) and '|' in lines[idx]:
            table_lines.append(lines[idx])
            idx += 1

        if len(table_lines) < 2:
            return idx

        # Парсим таблицу
        rows = []
        for line in table_lines:
            # Пропускаем разделительную строку (---|---|---)
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue

            # Разбиваем строку на ячейки
            cells = [cell.strip() for cell in line.split('|')]
            # Удаляем пустые ячейки в начале и конце
            cells = [c for c in cells if c or cells.index(c) not in [0, len(cells) - 1]]
            if cells:
                rows.append(cells)

        if not rows:
            return idx

        # Создаем таблицу в документе
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        self.add_table_borders(table)

        # Заполняем таблицу
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    # Очищаем ячейку и добавляем форматированный текст
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    self.add_formatted_text(paragraph, cell_text)
                    # Устанавливаем выравнивание по ширине для текста в ячейках
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    # Форматирование для заголовка таблицы (первая строка)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

        return idx

    def process_list(self, lines, start_idx):
        """Обрабатывает списки (маркированные и нумерованные)"""
        idx = start_idx
        list_items = []

        # Определяем тип списка
        first_line = lines[start_idx].strip()
        is_ordered = bool(re.match(r'^\d+\.', first_line))

        # Собираем элементы списка
        while idx < len(lines):
            line = lines[idx].strip()
            if not line:
                break

            # Проверяем маркированный список
            if re.match(r'^[-*+]\s', line):
                list_items.append(('unordered', line[2:]))
                idx += 1
            # Проверяем нумерованный список
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^\d+\.\s(.*)', line)
                list_items.append(('ordered', match.group(1)))
                idx += 1
            else:
                break

        # Добавляем элементы списка в документ
        for list_type, item_text in list_items:
            paragraph = self.doc.add_paragraph(style='List Bullet' if list_type == 'unordered' else 'List Number')
            self.add_formatted_text(paragraph, item_text)

        return idx

    def convert_file(self, input_path, output_path):
        """Конвертирует markdown файл в Word"""
        try:
            # Читаем файл
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Создаем документ
            self.create_document()

            # Разбиваем на строки
            lines = content.split('\n')

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # Пропускаем пустые строки
                if not stripped:
                    i += 1
                    continue

                # Заголовки
                if stripped.startswith('#'):
                    level = 0
                    while level < len(stripped) and stripped[level] == '#':
                        level += 1

                    title_text = stripped[level:].strip()
                    heading_style = f'Heading {min(level, 9)}'

                    paragraph = self.doc.add_paragraph(style=heading_style)
                    self.add_formatted_text(paragraph, title_text)
                    paragraph.paragraph_format.space_after = Pt(6)
                    i += 1

                # Горизонтальная линия
                elif stripped == '---' or stripped == '***' or stripped == '___':
                    self.doc.add_paragraph('_' * 50)
                    i += 1

                # Таблицы
                elif '|' in line:
                    i = self.process_table(lines, i)

                # Списки
                elif re.match(r'^[-*+]\s', stripped) or re.match(r'^\d+\.\s', stripped):
                    i = self.process_list(lines, i)

                # Блок кода (```)
                elif stripped.startswith('```'):
                    i += 1
                    code_lines = []
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1

                    if code_lines:
                        paragraph = self.doc.add_paragraph()
                        run = paragraph.add_run('\n'.join(code_lines))
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)

                    i += 1

                # Обычный текст
                else:
                    paragraph = self.doc.add_paragraph()
                    self.add_formatted_text(paragraph, line)
                    # Устанавливаем выравнивание по ширине для обычного текста
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    i += 1

            # Сохраняем документ
            self.doc.save(output_path)
            return True, "Успешно конвертировано"

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"


def convert_markdown_to_word(input_path, output_path=None):
    """
    Основная функция для конвертации Markdown в Word

    Args:
        input_path (str): Путь к входному .md файлу
        output_path (str, optional): Путь к выходному .docx файлу.
                                   Если не указан, создается автоматически.

    Returns:
        tuple: (success: bool, message: str)
    """
    # Если выходной путь не указан, создаем его на основе входного
    if output_path is None:
        input_path_obj = Path(input_path)
        output_path = input_path_obj.with_suffix('.docx')

    # Создаем конвертер и выполняем конвертацию
    converter = MarkdownToWordConverter()
    return converter.convert_file(input_path, output_path)


def convert_multiple_files(input_files, output_directory=None):
    """
    Конвертирует несколько файлов

    Args:
        input_files (list): Список путей к .md файлам
        output_directory (str, optional): Папка для сохранения.
                                        Если не указана, файлы сохраняются рядом с исходными.

    Returns:
        dict: Результаты конвертации {filename: (success, message)}
    """
    results = {}

    for input_file in input_files:
        try:
            # Определяем путь для сохранения
            if output_directory:
                filename = Path(input_file).stem + ".docx"
                output_path = os.path.join(output_directory, filename)
            else:
                output_path = None

            # Конвертируем файл
            success, message = convert_markdown_to_word(input_file, output_path)
            results[input_file] = (success, message)

            print(f"{'✓' if success else '✗'} {Path(input_file).name}: {message}")

        except Exception as e:
            error_msg = f"Ошибка: {str(e)}"
            results[input_file] = (False, error_msg)
            print(f"✗ {Path(input_file).name}: {error_msg}")

    return results


# Примеры использования
if __name__ == "__main__":
    # Пример 1: Конвертация одного файла
    input_file = "./generated_kp/gg.md"
    output_file = "./generated_kp/example.docx"

    success, message = convert_markdown_to_word(input_file, output_file)
    print(f"Результат: {message}")

    # Пример 2: Конвертация одного файла с автоматическим именем
    # success, message = convert_markdown_to_word("document.md")
    # print(f"Результат: {message}")
    #
    # # Пример 3: Конвертация нескольких файлов
    # files_to_convert = [
    #     "file1.md",
    #     "file2.md",
    #     "file3.md"
    # ]
    #
    # results = convert_multiple_files(files_to_convert, "output_folder")
    #
    # # Статистика
    # successful = sum(1 for result in results.values() if result[0])
    # print(f"\nКонвертация завершена. Успешно: {successful}/{len(files_to_convert)}")