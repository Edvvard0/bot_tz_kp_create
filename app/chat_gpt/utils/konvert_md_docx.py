import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    """Класс для конвертации Markdown в Word с точным форматированием"""

    def __init__(self):
        self.doc = None
        # Путь к логотипу - используем относительный путь от корня проекта
        self.logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "hacktaika.png")

    def create_document(self):
        """Создает новый документ Word с логотипом в левом верхнем углу"""
        self.doc = Document()

        # Устанавливаем стандартный стиль для основного текста (13pt)
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Onest'
        font.size = Pt(13)
        font.color.rgb = RGBColor(0, 0, 0)

    def add_header_with_logo(self, project_name):
        """Добавляет шапку с названием проекта"""
        try:
            # Добавляем название проекта
            if project_name:
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(project_name)
                run.bold = True
                run.font.size = Pt(23)
                run.font.name = 'Onest'
                para.paragraph_format.space_after = Pt(6)
        except Exception as e:
            print(f"⚠️ Не удалось добавить шапку: {e}")

    def add_project_info(self, project_name, creation_date):
        """Добавляет информацию о проекте под шапкой"""
        # Добавляем шапку с названием проекта
        self.add_header_with_logo(project_name)

        # Добавляем дату создания (выровнено по левому краю)
        if creation_date:
            date_paragraph = self.doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            date_run = date_paragraph.add_run(f"Дата создания: {creation_date}")
            date_run.italic = True
            date_run.font.size = Pt(10)
            date_run.font.name = 'Onest'
            date_paragraph.paragraph_format.space_after = Pt(12)

    def add_section_title(self, title, level=2):
        """Добавляет заголовок раздела"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run(title)
        run.bold = True

        # Устанавливаем размер шрифта в зависимости от уровня
        if level == 1:
            run.font.size = Pt(15)  # Для "План работы"
        elif level == 2:
            run.font.size = Pt(13)  # Для "Краткое описание проекта"
        else:
            run.font.size = Pt(11)

        run.font.name = 'Onest'
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.space_before = Pt(12)

    def add_table_borders(self, table):
        """Добавляет границы к таблице"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Создаем элемент границ
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def parse_inline_formatting(self, text):
        """Парсит встроенное форматирование (жирный, курсив)"""
        parts = []

        # Регулярное выражение для поиска форматирования
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|___.*?___|__.*?__|_.*?_)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Добавляем текст до совпадения
            if match.start() > last_end:
                parts.append((text[last_end:match.start()], {}))

            matched_text = match.group()
            formatting = {}
            clean_text = matched_text

            # Проверяем тип форматирования
            if matched_text.startswith('***') or matched_text.startswith('___'):
                formatting = {'bold': True, 'italic': True}
                clean_text = matched_text[3:-3]
            elif matched_text.startswith('**') or matched_text.startswith('__'):
                formatting = {'bold': True}
                clean_text = matched_text[2:-2]
            elif matched_text.startswith('*') or matched_text.startswith('_'):
                formatting = {'italic': True}
                clean_text = matched_text[1:-1]

            parts.append((clean_text, formatting))
            last_end = match.end()

        # Добавляем оставшийся текст
        if last_end < len(text):
            parts.append((text[last_end:], {}))

        return parts if parts else [(text, {})]

    def add_formatted_text(self, paragraph, text, font_size=13):
        """Добавляет текст с форматированием в параграф"""
        parts = self.parse_inline_formatting(text)

        for part_text, formatting in parts:
            run = paragraph.add_run(part_text)
            run.font.name = 'Onest'
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)

            if formatting.get('bold'):
                run.bold = True
            if formatting.get('italic'):
                run.italic = True

    def process_table(self, lines, start_idx):
        """Обрабатывает таблицу из markdown"""
        table_lines = []
        idx = start_idx

        # Собираем все строки таблицы
        while idx < len(lines) and '|' in lines[idx]:
            table_lines.append(lines[idx])
            idx += 1

        if len(table_lines) < 2:
            return idx

        # Парсим таблицу
        rows = []
        for line in table_lines:
            # Пропускаем разделительную строку (---|---|---)
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue

            # Разбиваем строку на ячейки
            cells = line.split('|')
            # Убираем пустые ячейки в начале и конце, а также делаем strip
            cells = [c.strip() for c in cells[1:-1] if c.strip()]
            if cells:
                rows.append(cells)

        if not rows:
            return idx

        # Создаем таблицу в документе
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        self.add_table_borders(table)

        # Заполняем таблицу
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    # Очищаем ячейку и добавляем форматированный текст
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    self.add_formatted_text(paragraph, cell_text)
                    # Устанавливаем выравнивание по левому краю для текста в ячейках
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Форматирование для заголовка таблицы (первая строка)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

        # Добавляем пустую строку после таблицы
        self.doc.add_paragraph()

        return idx

    def convert_file(self, input_path, output_path, project_name=None, creation_date=None):
        """Конвертирует markdown файл в Word с точным форматированием"""
        try:
            # Читаем файл
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Создаем документ с логотипом
            self.create_document()

            # Извлекаем название проекта из markdown, если не было передано
            if not project_name:
                # Ищем первый заголовок с форматом # **Проект: название**
                for line in content.split('\n'):
                    stripped = line.strip()
                    if stripped.startswith('# **Проект:'):
                        # Извлекаем название из формата # **Проект: Название**
                        # Убираем начало и конец
                        text = stripped.replace('# ', '').replace('**', '').replace('Проект:', '')
                        project_name = text.strip()
                        if project_name:
                            break
            
            # Если все еще нет названия, используем заглушку
            if not project_name:
                project_name = "Коммерческое предложение"

            # Добавляем информацию о проекте
            self.add_project_info(project_name, creation_date)

            # Разбиваем на строки
            lines = content.split('\n')

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # Пропускаем пустые строки
                if not stripped:
                    i += 1
                    continue

                # Пропускаем основной заголовок с названием проекта (уже добавили в шапке)
                if stripped.startswith('# **Проект:') or (stripped.startswith('# **') and 'Проект:' in stripped):
                    i += 1
                    continue

                # Заголовки разделов (## заголовок)
                if stripped.startswith('## '):
                    title_text = stripped[3:].strip()
                    # Убираем ** вокруг текста
                    title_text = re.sub(r'^\*\*(.*?)\*\*$', r'\1', title_text)
                    self.add_section_title(title_text, level=1)
                    i += 1
                    continue

                # Подзаголовки (### заголовок)
                elif stripped.startswith('### '):
                    title_text = stripped[4:].strip()
                    # Убираем ** вокруг текста
                    title_text = re.sub(r'^\*\*(.*?)\*\*$', r'\1', title_text)
                    self.add_section_title(title_text, level=2)
                    i += 1
                    continue

                # Горизонтальная линия
                elif stripped == '---' or stripped == '***' or stripped == '___':
                    self.doc.add_paragraph()
                    i += 1
                    continue

                # Таблицы
                elif '|' in line:
                    i = self.process_table(lines, i)
                    continue

                # Обычный текст - сохраняем форматирование
                else:
                    if stripped:
                        paragraph = self.doc.add_paragraph()
                        self.add_formatted_text(paragraph, stripped)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.space_after = Pt(6)
                    i += 1

            # Сохраняем документ
            self.doc.save(output_path)
            return True, "Успешно конвертировано"

        except Exception as e:
            return False, f"Ошибка при конвертации: {str(e)}"


def convert_markdown_to_word(input_path, output_path=None, project_name=None, creation_date=None):
    """
    Основная функция для конвертации Markdown в Word с точным форматированием
    """
    # Если выходной путь не указан, создаем его на основе входного
    if output_path is None:
        input_path_obj = Path(input_path)
        output_path = input_path_obj.with_suffix('.docx')

    # Создаем конвертер и выполняем конвертацию
    converter = MarkdownToWordConverter()
    return converter.convert_file(input_path, output_path, project_name, creation_date)


def convert_kp_markdown_to_word(input_path, output_path, project_name, creation_date):
    """
    Специальная функция для конвертации КП Markdown в Word с точным форматированием
    """
    return convert_markdown_to_word(input_path, output_path, project_name, creation_date)

def main():
    input_file = r"F:\my_projects\my_bot_tz_kp\app\chat_gpt\generated_kp\gg.md"
    # Проверяем существование входного файла
    if not os.path.exists(input_file):
        print(f"Ошибка: Файл '{input_file}' не найден")
        sys.exit(1)

    # Определяем выходной файл
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        # Создаем имя выходного файла на основе входного
        input_path = Path(input_file)
        output_file = input_path.with_suffix('.docx')

    print(f"Конвертация: {input_file} -> {output_file}")

    # Выполняем конвертацию
    success, message = convert_markdown_to_word(input_file, output_file)

    if success:
        print(f"✅ Успешно: {message}")
        print(f"📄 Файл сохранен: {output_file}")
    else:
        print(f"❌ Ошибка: {message}")
        sys.exit(1)


if __name__ == "__main__":
    main()